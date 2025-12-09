import streamlit as st
from docx import Document
from io import BytesIO
import os
from datetime import date
import time

# --- 1. CONFIGURAÇÃO ---
st.set_page_config(page_title="Sistema HOF - Cloud", layout="wide")

# --- 2. LOGIN ---
USUARIOS_PERMITIDOS = {
    "willians": "Re105763#",
    "paula": "Re121091"
}

def check_password():
    if st.session_state.get('password_correct', False): return True
    st.markdown("<h1 style='text-align: center;'>🔒 Acesso Restrito HOF</h1>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        usuario = st.text_input("Usuário")
        senha = st.text_input("Senha", type="password")
        if st.button("ENTRAR", type="primary", use_container_width=True):
            if usuario in USUARIOS_PERMITIDOS and USUARIOS_PERMITIDOS[usuario] == senha:
                st.session_state['password_correct'] = True
                st.session_state['usuario_atual'] = usuario
                st.rerun()
            else: st.error("❌ Acesso Negado")
    return False

if not check_password(): st.stop()

# --- 3. MAPEAMENTO ---
MAPA_ARQUIVOS = {
    "Toxina Botulínica": "toxina",
    "Preenchimento Facial": "preenchimento",
    "Bioestimulador": "bioestimulador",
    "Fios de Sustentação": "fios",
    "Lipo Mecânica de Papada": "lipomecanica",
    "Lipo Enzimática de Papada": "lipoenzimatica",
    "Bichectomia": "bichectomia",
    "Microagulhamento": "microagulhamento",
    "Peeling": "peeling"
}

# --- 4. FUNÇÕES MELHORADAS ---

def formatar_real(valor):
    return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def converter_numero_texto(dias):
    """Agora cobre todos os dias do mês!"""
    numeros = {
        0: "zero", 1: "um", 2: "dois", 3: "três", 4: "quatro", 5: "cinco", 
        6: "seis", 7: "sete", 8: "oito", 9: "nove", 10: "dez",
        11: "onze", 12: "doze", 13: "treze", 14: "quatorze", 15: "quinze",
        16: "dezesseis", 17: "dezessete", 18: "dezoito", 19: "dezenove", 20: "vinte",
        21: "vinte e um", 22: "vinte e dois", 23: "vinte e três", 24: "vinte e quatro",
        25: "vinte e cinco", 26: "vinte e seis", 27: "vinte e sete", 28: "vinte e oito",
        29: "vinte e nove", 30: "trinta", 31: "trinta e um"
    }
    return numeros.get(dias, str(dias))

def substituir_no_paragrafo(paragrafo, refs):
    """
    Função Cirúrgica: Tenta substituir mantendo a formatação original (Negrito, Tamanho, etc).
    """
    if not paragrafo.text:
        return

    for key, value in refs.items():
        if key in paragrafo.text:
            # Tentativa 1: Substituir dentro do "Run" (pedaço formatado) para manter o estilo
            substituiu_com_estilo = False
            for run in paragrafo.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)
                    substituiu_com_estilo = True
            
            # Tentativa 2: Se a etiqueta estiver "quebrada" pelo Word, faz a substituição bruta
            # (Isso corrige o erro de não substituir, mas pode perder formatação específica)
            if not substituiu_com_estilo:
                paragrafo.text = paragrafo.text.replace(key, value)

def preencher_template(caminho, dados):
    if not os.path.exists(caminho): return None
    doc = Document(caminho)
    
    val_cheio = formatar_real(dados.get('valor_cheio', 0))
    val_desc = formatar_real(dados.get('valor_desconto', 0))
    val_final = formatar_real(dados.get('valor_final', 0))
    data_hoje = date.today().strftime("%d/%m/%Y")
    
    cid_valor = dados.get('cid', "")
    texto_cid_final = f"CID: {cid_valor}" if cid_valor else ""
    
    refs = {
        "{{NOME_PACIENTE}}": dados.get('nome', ""),
        "{{RG_PACIENTE}}": dados.get('rg', ""),
        "{{CPF_PACIENTE}}": dados.get('cpf', ""),
        "{{CELULAR_PACIENTE}}": dados.get('celular', ""),
        "{{ENDERECO_PACIENTE}}": dados.get('endereco', ""),
        "{{DATA_HOJE}}": data_hoje,
        "{{DESCRIÇÃO_PROCEDIMENTOS}}": ", ".join(dados.get('procedimentos', [])),
        "{{VALOR_CHEIO}}": val_cheio,
        "{{VALOR_DESCONTO}}": val_desc,
        "{{VALOR_FINAL}}": val_final,
        "{{FORMA_PAGAMENTO}}": dados.get('pagamento', ""),
        "{{CLAUSULA_IMAGEM}}": dados.get('clausula_imagem', ""),
        "{{LISTA_MEDICAMENTOS}}": dados.get('texto_medicamentos', ""),
        "{{DIAS_NUMERO}}": str(dados.get('dias_afastamento', 0)),
        "{{DIAS_EXTENSO}}": dados.get('dias_extenso', ""),
        "{{CID}}": texto_cid_final
    }
    
    # 1. Substitui no texto corrido
    for p in doc.paragraphs:
        substituir_no_paragrafo(p, refs)
        
    # 2. Substitui dentro de TABELAS (Correção para Contratos/Fichas)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    substituir_no_paragrafo(p, refs)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 5. INTERFACE ---
with st.sidebar:
    st.success(f"🟢 Usuário: {st.session_state['usuario_atual']}")
    if st.button("Sair"):
        st.session_state['password_correct'] = False
        st.rerun()
    st.markdown("---")
    st.header("👤 Paciente")
    nome = st.text_input("Nome Completo")
    rg = st.text_input("RG")
    cpf = st.text_input("CPF")
    celular = st.text_input("Celular")
    endereco = st.text_area("Endereço")

st.title("💉 Sistema Integrado HOF")
st.markdown("---")

col1, col2 = st.columns(2)
with col1:
    procs = st.multiselect("Procedimentos", list(MAPA_ARQUIVOS.keys()))

opcoes_docs = [
    "Contrato de Serviço", "Orçamento", "Recibo de Pagamento",
    "Autorização Tratamento Estético", "Uso de Imagem",
    "Termos de Consentimento (Específicos)", "Cuidados Pós (Específicos)",
    "Prontuário", "Anamnese", "Receituário", "Atestado Médico"
]

with col2:
    docs = st.multiselect("Selecione os Documentos", opcoes_docs)

st.markdown("---")

valor_cheio, valor_desconto, valor_final = 0.0, 0.0, 0.0
pgto, dias, dias_extenso, cid = "", 0, "", ""
txt_clausula, txt_receita = "", ""

if docs:
    st.subheader("📝 Preenchimento")
    
    # Lógica Financeira
    if any(d in docs for d in ["Contrato de Serviço", "Recibo de Pagamento", "Orçamento"]):
        st.info("💰 Financeiro")
        c1, c2, c3 = st.columns(3)
        valor_cheio = c1.number_input("Valor Original (R$)", 0.0, step=50.0)
        valor_desconto = c2.number_input("Desconto (R$)", 0.0, step=50.0)
        valor_final = valor_cheio - valor_desconto
        c3.metric("Valor Final", f"R$ {formatar_real(valor_final)}")
        pgto = st.text_area("Forma de Pagamento")
        if valor_desconto > 0:
            txt_clausula = f"Desconto de imagem: R$ {formatar_real(valor_desconto)}."

    # Receita
    if "Receituário" in docs:
        st.info("💊 Receituário")
        if 'lista_meds' not in st.session_state: st.session_state.lista_meds = []
        c_rem1, c_rem2 = st.columns([3, 1])
        med = c_rem1.text_input("Remédio")
        if c_rem2.button("Add") and med: st.session_state.lista_meds.append(med)
        for i, m in enumerate(st.session_state.lista_meds):
            st.text(f"- {m}")
            txt_receita += f"{i+1}. {m}\n"
        if st.button("Limpar"): st.session_state.lista_meds = []

    # Atestado
    if "Atestado Médico" in docs:
        st.info("crm Atestado")
        dias = st.number_input("Dias", 1)
        dias_extenso = converter_numero_texto(dias)
        cid = st.text_input("CID (Opcional)")

    st.markdown("---")
    
    if st.button("GERAR DOCUMENTOS 📂", type="primary"):
        if not nome:
            st.error("⚠️ Preencha o Nome do Paciente!")
        else:
            dados = {
                'nome': nome, 'rg': rg, 'cpf': cpf, 'celular': celular, 'endereco': endereco,
                'procedimentos': procs, 'valor_cheio': valor_cheio, 'valor_desconto': valor_desconto, 
                'valor_final': valor_final, 'pagamento': pgto, 'clausula_imagem': txt_clausula,
                'texto_medicamentos': txt_receita, 'dias_afastamento': dias, 
                'dias_extenso': dias_extenso, 'cid': cid
            }
            st.success("Arquivos gerados! Baixe abaixo:")

            # Lista Geral
            gerais = {
                "Contrato de Serviço": "contrato_orofacial.docx",
                "Orçamento": "orcamento.docx",
                "Recibo de Pagamento": "recibo.docx",
                "Autorização Tratamento Estético": "autorizacao_estetico.docx",
                "Uso de Imagem": "autorizacao_imagem.docx",
                "Prontuário": "prontuario.docx",
                "Anamnese": "anamnese.docx",
                "Receituário": "receituario.docx",
                "Atestado Médico": "atestado.docx"
            }

            for doc_nome, arquivo_real in gerais.items():
                if doc_nome in docs:
                    arq = preencher_template(f"templates/{arquivo_real}", dados)
                    if arq: st.download_button(f"📥 {doc_nome}", arq, f"{doc_nome}_{nome}.docx")
                    else: st.warning(f"⚠️ Faltou: templates/{arquivo_real}")

            # Lista Específicos
            if "Termos de Consentimento (Específicos)" in docs:
                for proc in procs:
                    sufixo = MAPA_ARQUIVOS.get(proc)
                    arq = preencher_template(f"templates/termo_{sufixo}.docx", dados)
                    if arq: st.download_button(f"📥 Termo - {proc}", arq, f"Termo_{sufixo}.docx")
            
            if "Cuidados Pós (Específicos)" in docs:
                for proc in procs:
                    sufixo = MAPA_ARQUIVOS.get(proc)
                    arq = preencher_template(f"templates/cuidados_{sufixo}.docx", dados)
                    if arq: st.download_button(f"📥 Cuidados - {proc}", arq, f"Cuidados_{sufixo}.docx")
